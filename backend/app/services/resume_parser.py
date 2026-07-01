"""
Resume parsing service.

Validates files by magic bytes (not MIME or extension) before passing to parsers.
Raw file bytes are never persisted — only the extracted + cleaned text.
"""

import io
import re

from fastapi import HTTPException, UploadFile, status

from app.core.config import settings

# Magic byte signatures for supported formats
_PDF_MAGIC = b"%PDF"
_DOCX_MAGIC = b"PK"  # DOCX is a ZIP archive; ZIP starts with PK\x03\x04


async def parse_resume(file: UploadFile) -> tuple[str, str, str]:
    """
    Read, validate, and parse an uploaded resume.

    Returns:
        (raw_text, file_type, safe_filename)

    Raises:
        HTTPException 413 — file too large
        HTTPException 415 — unsupported format
        HTTPException 422 — file corrupt, encrypted, or contains no readable text
    """
    content = await file.read()

    if len(content) > settings.max_upload_size_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File exceeds the 10 MB limit.",
        )

    file_type = _detect_file_type(content)
    raw_text = _extract_text(content, file_type)
    raw_text = _clean_text(raw_text)

    if not raw_text.strip():
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="No readable text found. The file may be scanned, image-only, or encrypted.",
        )

    safe_filename = _sanitize_filename(file.filename or "resume", file_type)
    return raw_text, file_type, safe_filename


def _detect_file_type(content: bytes) -> str:
    if content[:4] == _PDF_MAGIC:
        return "pdf"
    if content[:2] == _DOCX_MAGIC:
        return "docx"
    raise HTTPException(
        status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        detail="Only PDF and DOCX files are accepted.",
    )


def _extract_text(content: bytes, file_type: str) -> str:
    if file_type == "pdf":
        return _extract_pdf(content)
    return _extract_docx(content)


def _extract_pdf(content: bytes) -> str:
    import pdfplumber

    parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Failed to parse PDF. The file may be corrupted or password-protected.",
        )
    return "\n\n".join(parts)


def _extract_docx(content: bytes) -> str:
    from docx import Document

    parts: list[str] = []
    try:
        doc = Document(io.BytesIO(content))
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
        # Tables often contain skills / education grids
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Failed to parse DOCX. The file may be corrupted.",
        )
    return "\n\n".join(parts)


def _clean_text(text: str) -> str:
    # Strip control characters except newlines and tabs
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    # Collapse 3+ consecutive blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse runs of spaces/tabs on the same line
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _sanitize_filename(filename: str, file_type: str) -> str:
    # Keep alphanumerics, spaces, hyphens, underscores, dots
    name = re.sub(r"[^\w\s\-.]", "", filename).strip()
    # Enforce correct extension
    if not name.lower().endswith(f".{file_type}"):
        name = f"{name}.{file_type}"
    return name[:255] or f"resume.{file_type}"
